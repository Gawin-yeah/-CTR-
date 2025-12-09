import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from openai import OpenAI
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook

# --- 页面配置 ---
st.set_page_config(page_title="CTR 终极稳定系统 (V56)", layout="wide")
st.title("🎯 首页卡片 CTR 终极稳定系统 (V56.0)")

# ==========================================
# 🧠 0. 状态记忆
# ==========================================
if 'persist_ex_a' not in st.session_state: st.session_state.persist_ex_a = []
if 'persist_ex_b' not in st.session_state: st.session_state.persist_ex_b = []
if 'persist_ex_dual' not in st.session_state: st.session_state.persist_ex_dual = []
if 'persist_in_a' not in st.session_state: st.session_state.persist_in_a = []
if 'persist_in_b' not in st.session_state: st.session_state.persist_in_b = []
if 'persist_in_dual' not in st.session_state: st.session_state.persist_in_dual = []

def update_ex_a(): st.session_state.persist_ex_a = st.session_state.k_ex_a
def update_ex_b(): st.session_state.persist_ex_b = st.session_state.k_ex_b
def update_ex_dual(): st.session_state.persist_ex_dual = st.session_state.k_ex_dual
def update_in_a(): st.session_state.persist_in_a = st.session_state.k_in_a
def update_in_b(): st.session_state.persist_in_b = st.session_state.k_in_b
def update_in_dual(): st.session_state.persist_in_dual = st.session_state.k_in_dual

# ==========================================
# 🛠️ 绘图与工具函数
# ==========================================
def plot_waterfall(df_waterfall, title):
    fig = go.Figure(go.Waterfall(
        name="20", orientation="v",
        measure=df_waterfall['measure'],
        x=df_waterfall['category'],
        textposition="outside",
        text=df_waterfall['text_val'],
        y=df_waterfall['value'],
        connector={"line": {"color": "rgb(63, 63, 63)"}},
        decreasing={"marker": {"color": "#EF553B"}},
        increasing={"marker": {"color": "#00CC96"}},
        totals={"marker": {"color": "#636EFA"}}
    ))
    fig.update_layout(title=title, showlegend=False, template="plotly_white", height=450)
    return fig

def plot_dual_axis(df, x_col, bar_col, line_col, title):
    fig = go.Figure()
    fig.add_trace(go.Bar(x=df[x_col], y=df[bar_col], name="总曝光", marker_color='#A9CCE3', opacity=0.6, yaxis='y1'))
    fig.add_trace(go.Scatter(x=df[x_col], y=df[line_col], name="CTR", mode='lines+markers', line=dict(color='#E74C3C', width=3), marker=dict(size=8), yaxis='y2'))
    fig.update_layout(title=title, xaxis_title="日期", yaxis=dict(title="曝光", side="left", showgrid=False), yaxis2=dict(title="CTR", side="right", overlaying="y", tickformat=".2%", showgrid=True), hovermode="x unified", legend=dict(orientation="h", y=1.1), template="plotly_white", height=400)
    return fig

def plot_bar_race(df, x_col, y_col, title):
    df[y_col] = df[y_col].astype(str)
    fig = px.bar(df, x=x_col, y=y_col, orientation='h', title=title, text_auto='.2%', color=x_col, color_continuous_scale='Blues')
    fig.update_layout(yaxis={'categoryorder':'total ascending', 'type': 'category'}, template="plotly_white", height=350, showlegend=False)
    return fig

def plot_pie(df, names, values, title):
    fig = px.pie(df, names=names, values=values, title=title, hole=0.4)
    fig.update_layout(template="plotly_white", height=350)
    return fig

def plot_paired_bar(df, category_col, val_a, val_b, title):
    df[category_col] = df[category_col].astype(str)
    df_melt = df.melt(id_vars=[category_col], value_vars=[val_a, val_b], var_name='时期', value_name='CTR')
    df_melt['时期'] = df_melt['时期'].map({val_a: '时期A', val_b: '时期B'})
    fig = px.bar(df_melt, y=category_col, x='CTR', color='时期', barmode='group', orientation='h', text_auto='.2%', title=title)
    fig.update_layout(yaxis={'categoryorder':'total ascending', 'type': 'category'}, xaxis_tickformat=".2%", height=500, legend=dict(orientation="h", y=1.1))
    return fig

def plot_impact_diverging(df, category_col, impact_col, title):
    df[category_col] = df[category_col].astype(str)
    df['Color'] = df[impact_col].apply(lambda x: '#E74C3C' if x >= 0 else '#2ECC71')
    fig = go.Figure(go.Bar(y=df[category_col], x=df[impact_col], orientation='h', marker=dict(color=df['Color']), text=df[impact_col], texttemplate='%{text:+.2%}', textposition='outside'))
    fig.update_layout(title=title, yaxis={'categoryorder':'total ascending', 'type': 'category'}, xaxis_tickformat=".2%", height=500)
    return fig

def generate_excel(dfs_dict):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        for sheet_name, df in dfs_dict.items():
            safe_name = sheet_name[:30]
            df.to_excel(writer, sheet_name=safe_name, index=False)
    return output.getvalue()

def generate_word_report(title, metrics, summary_text, tables_data):
    doc = Document()
    doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"生成时间: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_heading('一、核心大盘战报', level=1)
    p = doc.add_paragraph()
    for k, v in metrics.items(): p.add_run(f"{k}: {v}\n").bold = True
    doc.add_heading('二、深度归因与洞察', level=1)
    doc.add_paragraph(summary_text)
    for t_title, df in tables_data.items():
        if df.empty: continue
        doc.add_heading(f"三、{t_title}", level=1)
        t = doc.add_table(rows=1, cols=len(df.columns))
        t.style = 'Table Grid'
        for i, c in enumerate(df.columns): t.rows[0].cells[i].text = str(c)
        for _, r in df.iterrows():
            row = t.add_row()
            for i, v in enumerate(r):
                row.cells[i].text = f"{v:.2%}" if isinstance(v, float) and abs(v)<1 else str(v)
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

# ==========================================
# 🤖 AI 助手
# ==========================================
def init_ai_sidebar(context_data):
    st.sidebar.markdown("---")
    st.sidebar.header("🤖 AI 智能分析助手")
    with st.sidebar.expander("⚙️ 模型配置", expanded=False):
        api_key = st.text_input("API Key", type="password")
        base_url = st.text_input("Base URL", value="https://api.deepseek.com")
        model_name = st.text_input("Model Name", value="deepseek-chat")
    
    if "messages" not in st.session_state: st.session_state.messages = []
    for msg in st.session_state.messages:
        with st.sidebar.chat_message(msg["role"]): st.markdown(msg["content"])
    
    if prompt := st.sidebar.chat_input("问我..."):
        if not api_key: st.sidebar.error("请填入 API Key")
        else:
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.sidebar.chat_message("user"): st.markdown(prompt)
            with st.sidebar.chat_message("assistant"):
                msg_ph = st.empty()
                full_res = ""
                try:
                    client = OpenAI(api_key=api_key, base_url=base_url)
                    stream = client.chat.completions.create(
                        model=model_name,
                        messages=[{"role": "system", "content": f"基于数据回答：\n{context_data}"}] + [{"role": m["role"], "content": m["content"]} for m in st.session_state.messages],
                        stream=True,
                    )
                    for chunk in stream:
                        if chunk.choices[0].delta.content:
                            full_res += chunk.choices[0].delta.content
                            msg_ph.markdown(full_res + "▌")
                    msg_ph.markdown(full_res)
                    st.session_state.messages.append({"role": "assistant", "content": full_res})
                except Exception as e: st.error(str(e))

GLOBAL_DATA_CONTEXT = "暂无数据。"

# ==========================================
# 📂 数据接入
# ==========================================
st.sidebar.header("1. 数据接入")
manual_country = st.sidebar.text_input("✍️ 所属国家", value="US").upper()
read_visible_only = st.sidebar.checkbox("👁️ 只读取显示行 (剔除筛选隐藏)", value=False)

file_a = st.sidebar.file_uploader("上传主表格 (A)", type=["xlsx", "xls", "csv"], key="file_a")
sheet_name_a = 0
if file_a and file_a.name.endswith(('xlsx', 'xls')):
    try:
        xls = pd.ExcelFile(file_a)
        if len(xls.sheet_names) > 1: sheet_name_a = st.sidebar.selectbox(f"表A工作表:", xls.sheet_names, key="s_a")
    except: pass

file_b = st.sidebar.file_uploader("上传对比表格 (B)", type=["xlsx", "xls", "csv"], key="file_b")
sheet_name_b = 0
if file_b and file_b.name.endswith(('xlsx', 'xls')):
    try:
        xls = pd.ExcelFile(file_b)
        if len(xls.sheet_names) > 1: sheet_name_b = st.sidebar.selectbox(f"表B工作表:", xls.sheet_names, key="s_b")
    except: pass

st.sidebar.markdown("---")
min_exp_noise = st.sidebar.number_input("📉 单日最小曝光阈值 (去噪)", value=50, step=50)

def extract_start_date(s):
    s = str(s).strip()
    if "~" in s: return s.split("~")[0].strip()
    if "～" in s: return s.split("～")[0].strip()
    return s

@st.cache_data
def process_data(file, sheet_name=0, visible_only=False):
    try:
        if file.name.endswith('.csv'):
            raw_df = pd.read_csv(file)
        elif visible_only:
            wb = load_workbook(file, data_only=True, read_only=False)
            ws = wb.active if sheet_name == 0 else wb[sheet_name]
            data = []
            rows = ws.iter_rows(values_only=False)
            headers = None
            for row in rows:
                if ws.row_dimensions[row[0].row].hidden: continue
                vals = [c.value for c in row]
                if headers is None: headers = vals
                else: data.append(vals)
            raw_df = pd.DataFrame(data, columns=headers)
        else:
            raw_df = pd.read_excel(file, sheet_name=sheet_name)
            
        rename_map = {}
        for col in raw_df.columns:
            if "卡片" in col or "Card" in col: rename_map[col] = 'card_id'
            elif "坑位" in col or "Slot" in col: rename_map[col] = 'slot_id'
            elif "指标" in col: rename_map[col] = 'metric_name'
        df = raw_df.rename(columns=rename_map)
        
        required = ['card_id', 'metric_name']
        if not all(col in df.columns for col in required): return None
        
        if 'slot_id' not in df.columns: df['slot_id'] = 'Default'
        df['card_id'] = df['card_id'].astype(str)
        df['slot_id'] = df['slot_id'].astype(str)
        
        fixed = ['card_id', 'slot_id', 'metric_name', '合计', '均值', '总计', 'Total']
        dates = [c for c in df.columns if c not in fixed and "Unnamed" not in str(c)]
        if not dates: return None
        
        melted = df.melt(id_vars=['card_id', 'slot_id', 'metric_name'], value_vars=dates, var_name='raw_date', value_name='count')
        melted['date'] = pd.to_datetime(melted['raw_date'].apply(extract_start_date), errors='coerce').dt.date
        melted = melted.dropna(subset=['date'])
        melted['count'] = pd.to_numeric(melted['count'], errors='coerce').fillna(0)
        
        def get_type(t):
            if "曝光" in str(t): return "exposure_uv"
            if "点击" in str(t): return "click_uv"
            return None
        melted['type'] = melted['metric_name'].apply(get_type)
        melted = melted.dropna(subset=['type'])
        
        final = melted.pivot_table(index=['date', 'card_id', 'slot_id'], columns='type', values='count', aggfunc='sum').reset_index()
        for c in ['exposure_uv', 'click_uv']:
            if c not in final.columns: final[c] = 0
        return final
    except: return None

def filter_dataframe(df, min_exp):
    if df is None: return None
    return df[(df['exposure_uv'] >= min_exp) & (df['click_uv'] <= df['exposure_uv'])].copy()

# --- 4. 单文件视图 (V56 修复版) ---
def render_analysis_view(data, group_cols, view_name, unique_key_prefix):
    # 1. 计算核心指标 (基础表)
    period = data.groupby(group_cols).agg({'exposure_uv':'sum', 'click_uv':'sum'}).reset_index()
    period['加权CTR'] = period['click_uv']/period['exposure_uv']
    
    daily = data.groupby(group_cols + ['date']).agg({'exposure_uv':'sum', 'click_uv':'sum'}).reset_index()
    daily['daily_ctr'] = daily['click_uv']/daily['exposure_uv']
    
    arith = daily.groupby(group_cols)['daily_ctr'].mean().reset_index().rename(columns={'daily_ctr':'算术CTR'})
    
    # base_df 只包含汇总数据，不包含日期列，避免冲突
    base_df = pd.merge(period, arith, on=group_cols, how='left').sort_values('exposure_uv', ascending=False)
    
    # Label 处理
    display_base = base_df.copy()
    if 'slot_id' in group_cols: display_base['label'] = display_base['card_id'] + " (" + display_base['slot_id'] + ")"
    else: display_base['label'] = display_base['card_id']
    
    # 仪表盘
    with st.expander(f"📊 {view_name} - Leader 驾驶舱", expanded=True):
        c1, c2 = st.columns(2)
        with c1: st.plotly_chart(plot_pie(display_base.head(8), 'label', 'exposure_uv', "流量 Top 8"), use_container_width=True)
        with c2: 
            top_ctr = display_base[display_base['exposure_uv'] > data['exposure_uv'].mean()*0.1].head(10)
            if not top_ctr.empty:
                st.plotly_chart(plot_bar_race(top_ctr, '加权CTR', 'label', "高潜 Top 10"), use_container_width=True)
            else: st.info("数据不足")

    st.markdown("---")
    st.markdown(f"#### 📋 详细数据透视 ({view_name})")
    
    c_s1, c_s2 = st.columns([2, 1])
    with c_s1:
        search_vals = st.multiselect(f"🔍 搜索/筛选卡片", display_base['label'].unique(), key=f"search_{unique_key_prefix}")
    with c_s2:
        table_metric = st.radio("📊 表格展示每日指标:", ["每日 CTR", "每日 曝光", "每日 点击"], horizontal=True, key=f"tm_{unique_key_prefix}")
    
    # 动态计算 Pivot，避免列名冲突
    if table_metric == "每日 CTR":
        val_col, fmt_str = 'daily_ctr', '{:.2%}'
    elif table_metric == "每日 曝光":
        val_col, fmt_str = 'exposure_uv', '{:,.0f}'
    else:
        val_col, fmt_str = 'click_uv', '{:,.0f}'
        
    pivot = daily.pivot_table(index=group_cols, columns='date', values=val_col, aggfunc='sum' if val_col != 'daily_ctr' else 'mean')
    pivot.columns = [d.strftime('%m-%d') for d in pivot.columns]
    
    # 这里的 merge 是安全的，因为 base_df 没有日期列
    final_display = pd.merge(display_base, pivot, on=group_cols, how='left')
    
    if search_vals:
        final_display = final_display[final_display['label'].isin(search_vals)]
    
    cols = ['card_id', 'slot_id', '加权CTR', '算术CTR', 'exposure_uv', 'click_uv'] if 'slot_id' in group_cols else ['card_id', '加权CTR', '算术CTR', 'exposure_uv', 'click_uv']
    cols += [c for c in pivot.columns]
    
    fmt = {'加权CTR':'{:.2%}', '算术CTR':'{:.2%}', 'exposure_uv':'{:.0f}', 'click_uv':'{:.0f}'}
    for c in pivot.columns: fmt[c] = fmt_str
    
    st.dataframe(final_display[cols].style.format(fmt).background_gradient(subset=['加权CTR'], cmap='RdYlGn', axis=0), use_container_width=True, height=500)

    st.markdown("#### 📈 趋势下钻")
    default_trend = search_vals if search_vals else []
    sel = st.multiselect(f"选择对象画图", display_base['label'].unique(), default=default_trend, key=f"ms_{unique_key_prefix}")
    if sel:
        metric_choice = st.radio("趋势指标:", ["✨ CTR", "📊 曝光量", "👆 点击量"], horizontal=True, key=f"rd_{unique_key_prefix}")
        plot_df = daily.copy()
        if 'slot_id' in group_cols: plot_df['label'] = plot_df['card_id'] + " (" + plot_df['slot_id'] + ")"
        else: plot_df['label'] = plot_df['card_id']
        plot_df = plot_df[plot_df['label'].isin(sel)]
        
        if metric_choice == "✨ CTR": y_col, fmt_p = 'daily_ctr', ".2%"
        elif metric_choice == "📊 曝光量": y_col, fmt_p = 'exposure_uv', ".0f"
        else: y_col, fmt_p = 'click_uv', ".0f"
            
        st.plotly_chart(px.line(plot_df, x='date', y=y_col, color='label', markers=True).update_yaxes(tickformat=fmt_p), use_container_width=True)

def show_single_analysis(df, label="表格 A", is_secondary=False):
    if label == "表格 A":
        key_ex, key_in = "k_ex_a", "k_in_a"
        def_ex, def_in = st.session_state.persist_ex_a, st.session_state.persist_in_a
        cb_ex, cb_in = update_ex_a, update_in_a
    elif label == "表格 B":
        key_ex, key_in = "k_ex_b", "k_in_b"
        def_ex, def_in = st.session_state.persist_ex_b, st.session_state.persist_in_b
        cb_ex, cb_in = update_ex_b, update_in_b
    else: 
        key_ex, key_in = f"ex_{label}", f"in_{label}"
        def_ex, def_in = [], []
        cb_ex, cb_in = None, None

    st.markdown(f"## 🔎 {label} - 深度分析")
    
    if not is_secondary:
        if st.checkbox("⚔️ 开启表内对比", key=f"sw_{label}"):
            show_comparison_logic(df, df, f"{label}-A", f"{label}-B")
            return

    all_cards = sorted(df['card_id'].unique())
    valid_def_in = [x for x in def_in if x in all_cards]
    valid_def_ex = [x for x in def_ex if x in all_cards]

    col_f1, col_f2 = st.columns(2)
    with col_f1:
        include_list = st.multiselect("✅ 只看指定卡片", all_cards, default=valid_def_in, key=key_in, on_change=cb_in)
    with col_f2:
        exclude_list = st.multiselect("🚫 剔除指定卡片", all_cards, default=valid_def_ex, key=key_ex, on_change=cb_ex)
    
    sub_df_raw = df.copy()
    if include_list: sub_df_raw = sub_df_raw[sub_df_raw['card_id'].isin(include_list)]
    if exclude_list: sub_df_raw = sub_df_raw[~sub_df_raw['card_id'].isin(exclude_list)]
    
    min_d, max_d = sub_df_raw['date'].min(), sub_df_raw['date'].max()
    dr = st.date_input("选择周期", [min_d, max_d], key=f"dr_{label}")
    if len(dr) != 2: return
    
    sub = sub_df_raw[(sub_df_raw['date']>=dr[0]) & (sub_df_raw['date']<=dr[1])].copy()
    
    e_tot = sub['exposure_uv'].sum()
    c_tot = sub['click_uv'].sum()
    ctr_w = c_tot/e_tot if e_tot>0 else 0
    daily_g = sub.groupby('date').agg({'exposure_uv':'sum', 'click_uv':'sum'}).reset_index()
    daily_g['ctr'] = daily_g['click_uv']/daily_g['exposure_uv']
    
    st.markdown("### 🌍 全盘趋势驾驶舱")
    st.plotly_chart(plot_dual_axis(daily_g, 'date', 'exposure_uv', 'ctr', "全盘流量 vs 效率"), use_container_width=True)
    
    c1, c2, c3 = st.columns(3)
    c1.metric("总曝光", f"{e_tot:,.0f}")
    c2.metric("总点击", f"{c_tot:,.0f}")
    c3.metric("加权均值 CTR", f"{ctr_w:.2%}")
    
    if not is_secondary:
        global GLOBAL_DATA_CONTEXT
        GLOBAL_DATA_CONTEXT = f"单表:{label}, 剔除:{exclude_list}, CTR:{ctr_w:.2%}, 曝光:{e_tot}"
    
    st.divider()
    t1, t2 = st.tabs(["💳 视图:只看卡片", "📍 视图:细分坑位"])
    with t1: render_analysis_view(sub, ['card_id'], "卡片维度", label+"1")
    with t2: render_analysis_view(sub, ['card_id', 'slot_id'], "坑位维度", label+"2")
    
    st.divider()
    st.header("📥 导出中心")
    c_e1, c_e2 = st.columns(2)
    export_df = sub.groupby(['card_id', 'slot_id']).agg({'exposure_uv':'sum', 'click_uv':'sum'}).reset_index()
    export_df['weighted_ctr'] = export_df['click_uv'] / export_df['exposure_uv']
    export_df = export_df.sort_values('exposure_uv', ascending=False)
    top_5 = export_df.head(5).rename(columns={'card_id':'卡片ID', 'weighted_ctr':'CTR', 'exposure_uv':'曝光'})
    
    word_file = generate_word_report(f"报告-{manual_country}", {"周期": str(dr), "曝光": f"{e_tot:,.0f}", "CTR": f"{ctr_w:.2%}"}, "数据详见附表", {"Top5": top_5})
    excel_file = generate_excel({"聚合": export_df, "明细": sub})
    with c_e1: st.download_button("📄 Word 报告", word_file, f"Report_{label}.docx", key=f"bw_{label}")
    with c_e2: st.download_button("📊 Excel 数据", excel_file, f"Data_{label}.xlsx", key=f"be_{label}")

# --- 5. 双表对比 ---
def show_comparison_logic(d1_raw, d2_raw, la="A", lb="B"):
    st.markdown("### ⚙️ 对比配置")
    mode = st.radio("维度", ["💳 仅卡片", "📍 卡片+坑位"], horizontal=True, key=f"rd_{la}")
    cols = ['card_id'] if "仅" in mode else ['card_id', 'slot_id']
    
    all_cards = sorted(list(set(d1_raw['card_id'])|set(d2_raw['card_id'])))
    
    if la == "表格A": 
        key_ex, key_in = "k_ex_dual", "k_in_dual"
        def_ex, def_in = st.session_state.persist_ex_dual, st.session_state.persist_in_dual
        cb_ex, cb_in = update_ex_dual, update_in_dual
    else: 
        key_ex, key_in = f"ex_{la}", f"in_{la}"
        def_ex, def_in = [], []
        cb_ex, cb_in = None, None

    valid_def_in = [x for x in def_in if x in all_cards]
    valid_def_ex = [x for x in def_ex if x in all_cards]

    col_f1, col_f2 = st.columns(2)
    with col_f1:
        inc = st.multiselect("✅ 只看指定卡片", all_cards, default=valid_def_in, key=key_in, on_change=cb_in)
    with col_f2:
        excl = st.multiselect("🚫 剔除指定卡片", all_cards, default=valid_def_ex, key=key_ex, on_change=cb_ex)
    
    d1, d2 = d1_raw.copy(), d2_raw.copy()
    if inc:
        d1 = d1[d1['card_id'].isin(inc)]
        d2 = d2[d2['card_id'].isin(inc)]
    if excl:
        d1 = d1[~d1['card_id'].isin(excl)]
        d2 = d2[~d2['card_id'].isin(excl)]
    
    c1, c2 = st.columns(2)
    dr1 = c1.date_input(f"{la} 时间", [d1['date'].min(), d1['date'].max()], key=f"d1_{la}")
    dr2 = c2.date_input(f"{lb} 时间", [d2['date'].min(), d2['date'].max()], key=f"d2_{la}")
    
    if len(dr1)==2 and len(dr2)==2:
        d1f = d1[(d1['date']>=dr1[0])&(d1['date']<=dr1[1])]
        d2f = d2[(d2['date']>=dr2[0])&(d2['date']<=dr2[1])]
        
        s1 = d1f.groupby(cols).agg({'exposure_uv':'sum', 'click_uv':'sum'}).reset_index()
        s2 = d2f.groupby(cols).agg({'exposure_uv':'sum', 'click_uv':'sum'}).reset_index()
        
        tea, tca = s1['exposure_uv'].sum(), s1['click_uv'].sum()
        teb, tcb = s2['exposure_uv'].sum(), s2['click_uv'].sum()
        ctra, ctrb = (tca/tea if tea>0 else 0), (tcb/teb if teb>0 else 0)
        
        df_m = pd.merge(s1, s2, on=cols, how='outer', suffixes=('_A', '_B')).fillna(0)
        df_m['CTRA'] = df_m.apply(lambda r: r['click_uv_A']/r['exposure_uv_A'] if r['exposure_uv_A']>0 else 0, axis=1)
        df_m['CTRB'] = df_m.apply(lambda r: r['click_uv_B']/r['exposure_uv_B'] if r['exposure_uv_B']>0 else 0, axis=1)
        
        df_m['WA'] = df_m['exposure_uv_A']/tea if tea>0 else 0
        df_m['WB'] = df_m['exposure_uv_B']/teb if teb>0 else 0
        
        df_m['IsNew'] = df_m['exposure_uv_A'] == 0
        df_m['IsLost'] = df_m['exposure_uv_B'] == 0
        df_m['IsCommon'] = (~df_m['IsNew']) & (~df_m['IsLost'])
        
        rate_eff = df_m[df_m['IsCommon']].apply(lambda r: (r['CTRB']-r['CTRA'])*r['WA'], axis=1).sum()
        mix_eff = df_m[df_m['IsCommon']].apply(lambda r: (r['WB']-r['WA'])*r['CTRA'], axis=1).sum()
        new_eff = df_m[df_m['IsNew']].apply(lambda r: (r['CTRB']-ctra)*r['WB'], axis=1).sum()
        lost_eff = df_m[df_m['IsLost']].apply(lambda r: (ctra-r['CTRA'])*r['WA'], axis=1).sum()
        
        df_m['Contrib'] = (df_m['click_uv_B']/teb if teb>0 else 0) - (df_m['click_uv_A']/tea if tea>0 else 0)
        
        ctr_diff = ctrb - ctra
        wf_df = pd.DataFrame({
            "measure": ["absolute", "relative", "relative", "relative", "relative", "total"],
            "category": ["A (基准)", "存量表现", "流量结构", "新卡红利", "下架/其他", "B (当前)"],
            "value": [ctra, rate_eff, mix_eff, new_eff, ctrb-ctra-rate_eff-mix_eff-new_eff, None],
            "text_val": [f"{ctra:.2%}", f"{rate_eff:+.2%}", f"{mix_eff:+.2%}", f"{new_eff:+.2%}", "Diff", f"{ctrb:.2%}"]
        })
        
        conclusion = ""
        if ctr_diff > 0:
            if new_eff > abs(rate_eff) and rate_eff < 0:
                conclusion = "🚀 **新卡驱动型**：本周期 CTR 提升主要是由**新素材**驱动的。\n⚠️ **警惕**：存量老卡片表现疲软（存量表现为负），且流量分配效率可能下降。"
            elif rate_eff > 0 and new_eff > 0:
                conclusion = "🌟 **全面普涨**：存量卡片质量提升，且新卡表现优异，业务处于健康增长期。"
            else:
                conclusion = "📈 **稳步增长**：各项指标均为正向贡献。"
        else:
            conclusion = "📉 **大盘回落**：需关注负向贡献最大的因子。"

        k1, k2, k3, k4 = st.columns(4)
        k1.metric("CTR", f"{ctrb:.2%}", f"{ctrb-ctra:+.2%}")
        k2.metric("倍数", f"{ctrb/ctra:.2f}x" if ctra>0 else "∞")
        k3.metric("曝光", f"{teb:,.0f}", f"{(teb-tea)/tea:+.1%}" if tea>0 else "∞")
        k4.metric("点击", f"{tcb:,.0f}", f"{(tcb-tca)/tca:+.1%}" if tca>0 else "∞")
        
        c_w, c_t = st.columns([2, 1])
        with c_w: 
            st.plotly_chart(plot_waterfall(wf_df, "CTR 涨跌归因瀑布"), use_container_width=True)
            with st.expander("📖 读懂这张图 (名词解释)"):
                st.markdown("- **存量表现**: 老卡片自身 CTR 变化的影响。\n- **流量结构**: 流量分配变化带来的影响。\n- **新卡红利**: 新上架卡片带来的增量。")
        with c_t: 
            st.success(f"**🤖 智能诊断**：\n\n{conclusion}")

        st.divider()
        st.subheader("🔎 量效气泡图 (存量卡片)")
        valid_scatter = df_m[df_m['IsCommon']].copy()
        if not valid_scatter.empty:
            valid_scatter['ExpChg'] = (valid_scatter['exposure_uv_B'] - valid_scatter['exposure_uv_A']) / (valid_scatter['exposure_uv_A'] + 1)
            valid_scatter['CTRChg'] = valid_scatter['CTRB'] - valid_scatter['CTRA']
            valid_scatter['label'] = valid_scatter['card_id']
            fig = px.scatter(valid_scatter, x="ExpChg", y="CTRChg", hover_name="label", size="exposure_uv_B", color="Contrib", color_continuous_scale="RdYlGn", title="曝光变化 vs CTR变化 (右上角=量价齐升)")
            fig.add_hline(y=0, line_dash="dash"); fig.add_vline(x=0, line_dash="dash")
            fig.update_xaxes(tickformat=".0%"); fig.update_yaxes(tickformat=".2%")
            st.plotly_chart(fig, use_container_width=True)

        st.subheader("🏆 贡献度排行榜 (Contribution)")
        def get_stat_label(r):
            if r['IsNew']: return '🟢 New'
            if r['IsLost']: return '🔴 Lost'
            return '🔵 Common'
        df_m['Status'] = df_m.apply(get_stat_label, axis=1)
        
        c_top, c_bot = st.columns(2)
        with c_top:
            st.markdown("**🚀 正向拉动 Top 5**")
            st.dataframe(df_m.sort_values('Contrib', ascending=False).head(5)[[cols[0], 'Status', 'Contrib', 'CTRB']].style.format({'Contrib':'+{:.2%}', 'CTRB':'{:.2%}'}), hide_index=True)
        with c_bot:
            st.markdown("**📉 负向拖累 Top 5**")
            st.dataframe(df_m.sort_values('Contrib', ascending=True).head(5)[[cols[0], 'Status', 'Contrib', 'CTRB']].style.format({'Contrib':'{:.2%}', 'CTRB':'{:.2%}'}), hide_index=True)

        st.divider()
        st.subheader("📋 详细数据表")
        show_cols = ['Status'] + cols + ['Contrib', 'exposure_uv_A', 'exposure_uv_B', 'CTRA', 'CTRB']
        st.dataframe(df_m[show_cols].sort_values('Contrib', ascending=False).style.format({'Contrib':'{:.2%}', 'CTRA':'{:.2%}', 'CTRB':'{:.2%}', 'exposure_uv_A':'{:.0f}', 'exposure_uv_B':'{:.0f}'}).background_gradient(subset=['Contrib'], cmap='RdYlGn', vmin=-0.005, vmax=0.005), use_container_width=True)
        
        st.divider()
        c_e1, c_e2 = st.columns(2)
        word_file = generate_word_report(f"归因战报-{manual_country}", {"CTR变化": f"{ctra:.2%}->{ctrb:.2%}"}, conclusion, {"贡献榜": df_m.head(5)})
        excel_file = generate_excel({"归因明细": df_m})
        with c_e1: st.download_button("📄 Word", word_file, f"Report_{la}.docx", key=f"bw_{la}")
        with c_e2: st.download_button("📊 Excel", excel_file, f"Data_{la}.xlsx", key=f"be_{la}")

def show_comparison(df1, df2):
    show_comparison_logic(df1, df2)

# --- 主逻辑 ---
df_a_raw = None
if file_a: df_a_raw = process_data(file_a, sheet_name_a, visible_only=read_visible_only)
df_b_raw = None
if file_b: df_b_raw = process_data(file_b, sheet_name_b, visible_only=read_visible_only)

# 全局清洗
df_a = filter_dataframe(df_a_raw, min_exp_noise)
df_b = filter_dataframe(df_b_raw, min_exp_noise)

if df_a is not None:
    if df_b is not None:
        mode = st.radio("👇 模式", ["📄 单文件分析", "⚔️ 双表对比"], horizontal=True)
        st.divider()
        if mode == "📄 单文件分析":
            t1, t2 = st.tabs(["表格 A", "表格 B"])
            with t1: show_single_analysis(df_a, "表格 A")
            with t2: show_single_analysis(df_b, "表格 B", is_secondary=True)
        else:
            show_comparison(df_a, df_b)
    else:
        show_single_analysis(df_a, "表格 A")
else:
    st.info("👈 请在左侧上传 Excel 文件。")

if GLOBAL_DATA_CONTEXT != "暂无数据。":
    init_ai_sidebar(GLOBAL_DATA_CONTEXT)
